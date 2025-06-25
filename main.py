# --- IMPORTS ---
from flask import Flask, render_template, request, jsonify, session, redirect, url_for, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash # Para contraseñas seguras
from docx import Document
import io
import os
from functools import wraps

# --- CONFIGURACIÓN DE LA APLICACIÓN ---
app = Flask(__name__)
# IMPORTANTE: ¡Cambia esto en una aplicación real!
app.config['SECRET_KEY'] = 'your_very_secret_key_for_session_management_v2'

# --- NUEVA CONFIGURACIÓN DE LA BASE DE DATOS ---
# Formato: 'mysql+pymysql://<usuario>:<contraseña>@<host>/<nombre_db>'
# Usando los datos de tu ejemplo de Sequelize. 'root' sin contraseña.
app.config['SQLALCHEMY_DATABASE_URI'] = 'mysql+pymysql://root:@localhost/actas'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Inicializa la extensión de la base de datos
db = SQLAlchemy(app)


# --- NUEVO MODELO DE USUARIO (EQUIVALENTE A TU MODELO DE SEQUELIZE) ---
# Esto define la tabla 'user' en tu base de datos.
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False) # Guardamos un hash, no la contraseña
    role = db.Column(db.String(20), nullable=False, default='user')
    is_active = db.Column(db.Boolean, default=True, nullable=False)

    # Método para establecer la contraseña (crea el hash)
    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    # Método para verificar la contraseña
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

    def __repr__(self):
        return f'<User {self.username}>'

# --- TUS FUNCIONES EXISTENTES DE PROCESAMIENTO DE DOCUMENTOS (SIN CAMBIOS) ---
def reemplazar_en_runs(parrafos, campos):
    for parrafo in parrafos:
        texto_total = "".join(run.text for run in parrafo.runs)
        original_texto_total = texto_total
        for clave, valor in campos.items():
            if clave in texto_total:
                texto_total = texto_total.replace(clave, valor)
        if texto_total != original_texto_total:
            for i in range(len(parrafo.runs)):
                parrafo.runs[i].text = ""
            if parrafo.runs:
                parrafo.runs[0].text = texto_total
            else:
                parrafo.add_run(texto_total)

def reemplazar_en_documento(doc, campos):
    reemplazar_en_runs(doc.paragraphs, campos)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                reemplazar_en_runs(celda.paragraphs, campos)

# --- DECORADORES DE PROTECCIÓN DE RUTAS (SIN CAMBIOS EN SU LÓGICA INTERNA) ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            if request.endpoint and (request.endpoint.startswith('admin_') or request.is_json):
                 return jsonify(success=False, message="Autenticación requerida."), 401
            return redirect(url_for('index', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if session.get('role') != 'admin':
            if request.endpoint and (request.endpoint.startswith('admin_') or request.is_json):
                return jsonify(success=False, message="Acceso de administrador requerido."), 403
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

# --- RUTA PRINCIPAL (SIN CAMBIOS) ---
@app.route('/')
def index():
    if 'username' in session:
        return render_template('index.html', logged_in_user=session.get('username'), user_role=session.get('role'))
    return render_template('index.html')

# --- RUTAS DE LOGIN/LOGOUT (MODIFICADAS PARA USAR LA BASE DE DATOS) ---
@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    # Busca al usuario en la base de datos
    user = User.query.filter_by(username=username).first()

    # Verifica si el usuario existe y si la contraseña es correcta (usando el hash)
    if user and user.check_password(password):
        if user.is_active:
            session['username'] = user.username
            session['role'] = user.role
            return jsonify({'success': True, 'role': session['role'], 'username': session['username']})
        else:
            return jsonify({'success': False, 'message': 'Cuenta desactivada.'})
    return jsonify({'success': False, 'message': 'Credenciales Incorrectas'})

@app.route('/logout')
@login_required
def logout():
    session.pop('username', None)
    session.pop('role', None)
    return redirect(url_for('index'))

# --- RUTAS DE ADMINISTRACIÓN (MODIFICADAS PARA USAR LA BASE DE DATOS) ---
@app.route('/admin/users', methods=['GET'])
@admin_required
def get_admin_users():
    # Consulta todos los usuarios de la base de datos
    users_from_db = User.query.all()
    # Prepara los datos para enviar como JSON (sin incluir el hash de la contraseña)
    users_safe = [
        {'username': u.username, 'role': u.role, 'is_active': u.is_active}
        for u in users_from_db
    ]
    return jsonify({'success': True, 'users': users_safe})

@app.route('/admin/users/add', methods=['POST'])
@admin_required
def add_admin_user():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')
    role = data.get('role', 'user')

    if not username or not password:
        return jsonify({'success': False, 'message': 'Usuario y contraseña son requeridos.'})

    # Verifica si el usuario ya existe en la base de datos
    if User.query.filter_by(username=username).first():
        return jsonify({'success': False, 'message': 'El nombre de usuario ya existe.'})

    # Crea una nueva instancia del modelo User
    new_user = User(username=username, role=role, is_active=True)
    new_user.set_password(password) # Hashea y guarda la contraseña

    # Añade y guarda el nuevo usuario en la base de datos
    db.session.add(new_user)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Usuario añadido correctamente.'})

@app.route('/admin/users/update', methods=['POST'])
@admin_required
def update_admin_user():
    data = request.get_json()
    username_to_update = data.get('username')
    is_active = data.get('is_active')
    new_password = data.get('new_password')
    new_role = data.get('new_role')

    # Busca al usuario que se quiere actualizar
    user_to_update = User.query.filter_by(username=username_to_update).first()
    
    if not user_to_update:
        return jsonify({'success': False, 'message': 'Usuario no encontrado.'})

    # Actualiza los campos si se proporcionaron
    if is_active is not None:
        user_to_update.is_active = is_active
    if new_role:
        user_to_update.role = new_role
    if new_password:
        user_to_update.set_password(new_password) # Actualiza la contraseña hasheada

    db.session.commit() # Guarda los cambios en la base de datos
    return jsonify({'success': True, 'message': 'Usuario actualizado.'})

@app.route('/admin/users/delete', methods=['POST'])
@admin_required
def delete_admin_user():
    data = request.get_json()
    username_to_delete = data.get('username')

    if username_to_delete == session.get('username'):
         return jsonify({'success': False, 'message': 'No puedes eliminar tu propia cuenta de administrador activa.'})

    # Busca al usuario a eliminar
    user_to_delete = User.query.filter_by(username=username_to_delete).first()

    if not user_to_delete:
        return jsonify({'success': False, 'message': 'Usuario no encontrado.'})
    
    # Elimina al usuario de la base de datos
    db.session.delete(user_to_delete)
    db.session.commit()

    return jsonify({'success': True, 'message': 'Usuario eliminado.'})


# --- RUTA DE GENERACIÓN DE ACTA (SIN CAMBIOS) ---
@app.route('/generar', methods=['POST'])
@login_required
def generar():
    campos = {}
    for clave, valor in request.form.items():
        if clave.endswith("_N") and not valor.strip():
            campos[f"{{{{{clave}}}}}"] = "0" 
        elif clave.endswith("_LI") and not valor.strip():
             campos[f"{{{{{clave}}}}}"] = "(vacío)"
        else:
            campos[f"{{{{{clave}}}}}"] = valor

    for i in range(1, 10):
        if f"{{{{ITEM{i}_N}}}}" not in campos:
            campos[f"{{{{ITEM{i}_N}}}}"] = "0"
        if f"{{{{ITEM{i}_LI}}}}" not in campos:
            campos[f"{{{{ITEM{i}_LI}}}}"] = "(no provisto)"
    
    if f"{{{{TOTAL_N}}}}" not in campos:
        campos[f"{{{{TOTAL_N}}}}"] = "0.00"
    if f"{{{{TOTAL_LI}}}}" not in campos:
        campos[f"{{{{TOTAL_LI}}}}"] = "(no calculado)"

    try:
        doc = Document("plantilla_acta.docx")
    except Exception as e:
        print(f"Error loading plantilla_acta.docx: {e}")
        return "Error: No se pudo cargar la plantilla del documento.", 500

    reemplazar_en_documento(doc, campos)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    postulante_name = request.form.get('POSTULANTE', 'documento_acta')
    safe_postulante_name = "".join(c for c in postulante_name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    filename = f"Acta_{safe_postulante_name.replace(' ', '_')}.docx"
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- INICIALIZACIÓN DE LA APLICACIÓN ---
if __name__ == '__main__':
    # Este bloque se asegura de que las tablas de la base de datos se creen
    # y crea un usuario 'admin' por defecto la primera vez que se ejecuta.
    with app.app_context():
        db.create_all() # Crea las tablas definidas en los modelos (ej. User) si no existen
        
        # Verifica si ya existe algún usuario en la base de datos
        if not User.query.first():
            print("Base de datos de usuarios vacía. Creando usuario admin por defecto...")
            admin_user = User(username='admin', role='admin', is_active=True)
            admin_user.set_password('changeme') # Contraseña por defecto
            db.session.add(admin_user)
            db.session.commit()
            print("Usuario 'admin' con contraseña 'changeme' creado. ¡Cámbiela desde el panel de administrador!")
    
    app.run(debug=True)
